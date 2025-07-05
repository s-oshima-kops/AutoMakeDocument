# -*- coding: utf-8 -*-
"""
出力設定ウィジェット
"""

from typing import Optional, Dict, Any, List
from pathlib import Path
from datetime import datetime
from PySide6.QtWidgets import (QWidget, QVBoxLayout, QHBoxLayout, QFormLayout,
                             QLabel, QGroupBox, QComboBox, QLineEdit, QPushButton,
                             QFileDialog, QTextEdit, QCheckBox, QSpinBox,
                             QTabWidget, QTableWidget, QTableWidgetItem,
                             QProgressBar, QListWidget, QListWidgetItem,
                             QMessageBox, QSplitter, QScrollArea)
from PySide6.QtCore import Qt, QThread, Signal, QTimer
from PySide6.QtGui import QFont

from core.template_engine import TemplateEngine
from utils.config_manager import ConfigManager
from utils.logger import Logger
from utils.file_utils import FileUtils

class OutputWorker(QThread):
    """出力処理ワーカー"""
    
    progress_updated = Signal(int)
    output_completed = Signal(str)  # file_path
    error_occurred = Signal(str)
    
    def __init__(self, template_engine: TemplateEngine, summary_data: Dict[str, Any],
                 template_id: str, output_config: Dict[str, Any], file_path: str):
        super().__init__()
        self.template_engine = template_engine
        self.summary_data = summary_data
        self.template_id = template_id
        self.output_config = output_config
        self.file_path = file_path
        
    def run(self):
        """出力処理を実行"""
        try:
            # 進捗更新
            self.progress_updated.emit(20)
            
            # テンプレートを適用
            template_result = self.template_engine.apply_template(
                self.template_id,
                self.summary_data,
                self.output_config
            )
            
            # 結果を文字列形式に変換
            formatted_content = self.template_engine.format_output(template_result, "text")
            
            self.progress_updated.emit(60)
            
            # ファイル出力
            self.write_output_file(formatted_content)
            
            self.progress_updated.emit(100)
            
            # 完了を通知
            self.output_completed.emit(self.file_path)
            
        except Exception as e:
            self.error_occurred.emit(str(e))
            
    def write_output_file(self, content: str):
        """出力ファイルを書き込み"""
        output_format = self.output_config.get('format', 'txt')
        
        if output_format == 'txt':
            self.write_text_file(content)
        elif output_format == 'csv':
            self.write_csv_file(content)
        elif output_format == 'xlsx':
            self.write_excel_file(content)
        elif output_format == 'docx':
            self.write_word_file(content)
        else:
            raise ValueError(f"サポートされていない出力形式: {output_format}")
            
    def write_text_file(self, content: str):
        """テキストファイルを書き込み"""
        with open(self.file_path, 'w', encoding='utf-8') as f:
            f.write(content)
            
    def write_csv_file(self, content: str):
        """CSVファイルを書き込み"""
        import csv
        
        # コンテンツを解析してCSV形式に変換
        lines = content.split('\n')
        csv_data = []
        
        current_section = ""
        for line in lines:
            line = line.strip()
            if line.startswith('■') or line.startswith('【'):
                current_section = line
            elif line and not line.startswith('・'):
                csv_data.append([current_section, line])
                
        with open(self.file_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(['セクション', '内容'])
            writer.writerows(csv_data)
            
    def write_excel_file(self, content: str):
        """Excelファイルを書き込み"""
        import openpyxl
        
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "要約レポート"
        
        # ヘッダー
        ws['A1'] = "要約レポート"
        ws['A1'].font = openpyxl.styles.Font(size=16, bold=True)
        
        # コンテンツを行ごとに書き込み
        lines = content.split('\n')
        row = 3
        
        for line in lines:
            if line.strip():
                ws[f'A{row}'] = line
                if line.startswith('■') or line.startswith('【'):
                    ws[f'A{row}'].font = openpyxl.styles.Font(bold=True)
                row += 1
                
        wb.save(self.file_path)
        
    def write_word_file(self, content: str):
        """Wordファイルを書き込み"""
        from docx import Document
        
        doc = Document()
        
        # タイトル
        title = doc.add_heading('要約レポート', level=1)
        
        # コンテンツを段落ごとに追加
        lines = content.split('\n')
        current_paragraph = []
        
        for line in lines:
            line = line.strip()
            if not line:
                if current_paragraph:
                    doc.add_paragraph('\n'.join(current_paragraph))
                    current_paragraph = []
            elif line.startswith('■'):
                if current_paragraph:
                    doc.add_paragraph('\n'.join(current_paragraph))
                    current_paragraph = []
                doc.add_heading(line[1:].strip(), level=2)
            elif line.startswith('【'):
                if current_paragraph:
                    doc.add_paragraph('\n'.join(current_paragraph))
                    current_paragraph = []
                doc.add_heading(line[1:-1].strip(), level=3)
            else:
                current_paragraph.append(line)
                
        if current_paragraph:
            doc.add_paragraph('\n'.join(current_paragraph))
            
        doc.save(self.file_path)

class OutputPreviewWidget(QWidget):
    """出力プレビューウィジェット"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setup_ui()
        
    def setup_ui(self):
        """UI設定"""
        layout = QVBoxLayout(self)
        
        # プレビューエリア
        self.preview_text = QTextEdit()
        self.preview_text.setReadOnly(True)
        self.preview_text.setFont(QFont("Consolas", 9))
        layout.addWidget(self.preview_text)
        
        # 文字数表示
        self.char_count_label = QLabel("文字数: 0")
        layout.addWidget(self.char_count_label)
        
    def set_preview_content(self, content: str):
        """プレビューコンテンツを設定"""
        self.preview_text.setPlainText(content)
        self.char_count_label.setText(f"文字数: {len(content)}")

class OutputHistoryWidget(QWidget):
    """出力履歴ウィジェット"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setup_ui()
        
    def setup_ui(self):
        """UI設定"""
        layout = QVBoxLayout(self)
        
        # 履歴リスト
        self.history_list = QListWidget()
        layout.addWidget(self.history_list)
        
        # 操作ボタン
        button_layout = QHBoxLayout()
        
        self.open_button = QPushButton("開く")
        self.open_button.clicked.connect(self.open_selected_file)
        button_layout.addWidget(self.open_button)
        
        self.delete_button = QPushButton("削除")
        self.delete_button.clicked.connect(self.delete_selected_file)
        button_layout.addWidget(self.delete_button)
        
        button_layout.addStretch()
        
        self.clear_button = QPushButton("履歴クリア")
        self.clear_button.clicked.connect(self.clear_history)
        button_layout.addWidget(self.clear_button)
        
        layout.addLayout(button_layout)
        
    def add_history_item(self, file_path: str, timestamp: datetime):
        """履歴アイテムを追加"""
        item_text = f"{timestamp.strftime('%Y-%m-%d %H:%M:%S')} - {Path(file_path).name}"
        item = QListWidgetItem(item_text)
        item.setData(Qt.UserRole, file_path)
        self.history_list.addItem(item)
        
    def open_selected_file(self):
        """選択されたファイルを開く"""
        current_item = self.history_list.currentItem()
        if current_item:
            file_path = current_item.data(Qt.UserRole)
            import os
            os.startfile(file_path)
            
    def delete_selected_file(self):
        """選択されたファイルを削除"""
        current_item = self.history_list.currentItem()
        if current_item:
            file_path = current_item.data(Qt.UserRole)
            try:
                Path(file_path).unlink()
                self.history_list.takeItem(self.history_list.row(current_item))
            except Exception as e:
                QMessageBox.warning(self, "エラー", f"ファイル削除エラー: {e}")
                
    def clear_history(self):
        """履歴をクリア"""
        self.history_list.clear()

class OutputConfigWidget(QWidget):
    """出力設定ウィジェット"""
    
    # シグナル定義
    output_completed = Signal(str)  # file_path
    
    def __init__(self, template_engine: TemplateEngine, config_manager: ConfigManager,
                 logger: Logger, app_dir: Path, parent=None):
        super().__init__(parent)
        
        self.template_engine = template_engine
        self.config_manager = config_manager
        self.logger = logger
        self.app_dir = app_dir
        
        self.current_summary = None
        self.current_template_id = None
        self.worker = None
        
        self.setup_ui()
        self.load_settings()
        
    def setup_ui(self):
        """UI設定"""
        layout = QVBoxLayout(self)
        
        # ヘッダー
        header_label = QLabel("出力設定")
        header_label.setStyleSheet("font-size: 16px; font-weight: bold; margin: 10px;")
        layout.addWidget(header_label)
        
        # メインスプリッター
        main_splitter = QSplitter(Qt.Horizontal)
        layout.addWidget(main_splitter)
        
        # 左側: 設定パネル
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        
        # 基本設定
        basic_group = QGroupBox("基本設定")
        basic_layout = QFormLayout(basic_group)
        
        # 出力形式
        self.format_combo = QComboBox()
        self.format_combo.addItems(["txt", "csv", "xlsx", "docx"])
        self.format_combo.currentTextChanged.connect(self.on_format_changed)
        basic_layout.addRow("出力形式:", self.format_combo)
        
        # 出力先フォルダ
        output_folder_layout = QHBoxLayout()
        self.output_folder_edit = QLineEdit()
        self.output_folder_edit.setText(str(self.app_dir / "output"))
        output_folder_layout.addWidget(self.output_folder_edit)
        
        self.browse_folder_button = QPushButton("参照")
        self.browse_folder_button.clicked.connect(self.browse_output_folder)
        output_folder_layout.addWidget(self.browse_folder_button)
        
        basic_layout.addRow("出力先フォルダ:", output_folder_layout)
        
        # ファイル名
        self.filename_edit = QLineEdit()
        self.filename_edit.setText(self.generate_default_filename())
        basic_layout.addRow("ファイル名:", self.filename_edit)
        
        left_layout.addWidget(basic_group)
        
        # テンプレート設定
        template_group = QGroupBox("テンプレート設定")
        template_layout = QFormLayout(template_group)
        
        # テンプレート選択
        self.template_combo = QComboBox()
        self.load_templates()
        template_layout.addRow("テンプレート:", self.template_combo)
        
        # カスタマイズ設定
        self.include_stats_check = QCheckBox("統計情報を含める")
        self.include_stats_check.setChecked(True)
        template_layout.addRow("", self.include_stats_check)
        
        self.include_keywords_check = QCheckBox("キーワードを含める")
        self.include_keywords_check.setChecked(True)
        template_layout.addRow("", self.include_keywords_check)
        
        self.include_original_check = QCheckBox("元テキストを含める")
        self.include_original_check.setChecked(False)
        template_layout.addRow("", self.include_original_check)
        
        left_layout.addWidget(template_group)
        
        # 出力実行
        execute_group = QGroupBox("実行")
        execute_layout = QVBoxLayout(execute_group)
        
        # プレビューボタン
        self.preview_button = QPushButton("プレビュー")
        self.preview_button.clicked.connect(self.preview_output)
        execute_layout.addWidget(self.preview_button)
        
        # 出力実行ボタン
        self.execute_button = QPushButton("出力実行")
        self.execute_button.clicked.connect(self.execute_output)
        execute_layout.addWidget(self.execute_button)
        
        # 進捗バー
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        execute_layout.addWidget(self.progress_bar)
        
        left_layout.addWidget(execute_group)
        
        left_layout.addStretch()
        main_splitter.addWidget(left_widget)
        
        # 右側: プレビューと履歴
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        
        # タブウィジェット
        self.right_tabs = QTabWidget()
        
        # プレビュータブ
        self.preview_widget = OutputPreviewWidget()
        self.right_tabs.addTab(self.preview_widget, "プレビュー")
        
        # 履歴タブ
        self.history_widget = OutputHistoryWidget()
        self.right_tabs.addTab(self.history_widget, "出力履歴")
        
        right_layout.addWidget(self.right_tabs)
        
        main_splitter.addWidget(right_widget)
        
        # 分割比率設定
        main_splitter.setSizes([400, 500])
        
    def load_templates(self):
        """テンプレートを読み込み"""
        try:
            templates = self.template_engine.get_available_templates()
            self.template_combo.clear()
            
            for template_info in templates:
                template_id = template_info.get('id', '')
                display_name = template_info.get('name', template_id)
                self.template_combo.addItem(display_name, template_id)
                
        except Exception as e:
            self.logger.log_error(f"テンプレート読み込みエラー: {e}")
            
    def generate_default_filename(self) -> str:
        """デフォルトファイル名を生成"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"summary_{timestamp}"
        
    def on_format_changed(self, format_type: str):
        """出力形式変更時"""
        # ファイル名の拡張子を更新
        filename = self.filename_edit.text()
        if '.' in filename:
            filename = filename.rsplit('.', 1)[0]
        self.filename_edit.setText(f"{filename}.{format_type}")
        
    def browse_output_folder(self):
        """出力フォルダを選択"""
        folder = QFileDialog.getExistingDirectory(
            self, "出力フォルダを選択", self.output_folder_edit.text()
        )
        if folder:
            self.output_folder_edit.setText(folder)
            
    def preview_output(self):
        """出力をプレビュー"""
        if not self.current_summary:
            QMessageBox.warning(self, "警告", "要約データがありません。")
            return
            
        try:
            # 設定を取得
            config = self.get_output_config()
            template_id = self.template_combo.currentData()
            
            # テンプレートを適用
            template_result = self.template_engine.apply_template(
                template_id, self.current_summary, config
            )
            
            # 結果を文字列形式に変換
            content = self.template_engine.format_output(template_result, "text")
            
            # プレビューに表示
            self.preview_widget.set_preview_content(content)
            self.right_tabs.setCurrentIndex(0)  # プレビュータブを表示
            
        except Exception as e:
            self.logger.log_error(f"プレビュー生成エラー: {e}")
            QMessageBox.critical(self, "エラー", f"プレビュー生成エラー: {e}")
            
    def execute_output(self):
        """出力を実行"""
        if not self.current_summary:
            QMessageBox.warning(self, "警告", "要約データがありません。")
            return
            
        try:
            # 出力ファイルパスを構築
            output_folder = Path(self.output_folder_edit.text())
            filename = self.filename_edit.text()
            file_path = output_folder / filename
            
            # フォルダが存在しない場合は作成
            output_folder.mkdir(parents=True, exist_ok=True)
            
            # 設定を取得
            config = self.get_output_config()
            template_id = self.template_combo.currentData()
            
            # ワーカーを開始
            self.start_output_worker(template_id, config, str(file_path))
            
        except Exception as e:
            self.logger.log_error(f"出力実行エラー: {e}")
            QMessageBox.critical(self, "エラー", f"出力実行エラー: {e}")
            
    def start_output_worker(self, template_id: str, config: Dict[str, Any], file_path: str):
        """出力ワーカーを開始"""
        # UI状態を更新
        self.execute_button.setEnabled(False)
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        
        # ワーカーを作成
        self.worker = OutputWorker(
            self.template_engine,
            self.current_summary,
            template_id,
            config,
            file_path
        )
        
        # シグナル接続
        self.worker.progress_updated.connect(self.progress_bar.setValue)
        self.worker.output_completed.connect(self.on_output_completed)
        self.worker.error_occurred.connect(self.on_output_error)
        
        # ワーカー開始
        self.worker.start()
        
    def on_output_completed(self, file_path: str):
        """出力完了時"""
        # UI状態を復元
        self.execute_button.setEnabled(True)
        self.progress_bar.setVisible(False)
        
        # 履歴に追加
        self.history_widget.add_history_item(file_path, datetime.now())
        
        # 完了メッセージ
        QMessageBox.information(self, "完了", f"出力が完了しました。\n{file_path}")
        
        # シグナルを発行
        self.output_completed.emit(file_path)
        
        # ログに記録
        self.logger.log_operation(f"出力完了: {file_path}")
        
        # 次のファイル名を生成
        self.filename_edit.setText(self.generate_default_filename())
        
    def on_output_error(self, error_message: str):
        """出力エラー時"""
        # UI状態を復元
        self.execute_button.setEnabled(True)
        self.progress_bar.setVisible(False)
        
        # エラーメッセージ
        QMessageBox.critical(self, "エラー", f"出力エラー: {error_message}")
        
        # ログに記録
        self.logger.log_error(f"出力エラー: {error_message}")
        
    def get_output_config(self) -> Dict[str, Any]:
        """出力設定を取得"""
        return {
            'format': self.format_combo.currentText(),
            'include_stats': self.include_stats_check.isChecked(),
            'include_keywords': self.include_keywords_check.isChecked(),
            'include_original': self.include_original_check.isChecked()
        }
        
    def set_summary_data(self, summary_data: Dict[str, Any]):
        """要約データを設定"""
        self.current_summary = summary_data
        
    def set_template_id(self, template_id: str):
        """テンプレートIDを設定"""
        self.current_template_id = template_id
        
        # コンボボックスで選択
        for i in range(self.template_combo.count()):
            if self.template_combo.itemData(i) == template_id:
                self.template_combo.setCurrentIndex(i)
                break
                
    def load_settings(self):
        """設定を読み込み"""
        try:
            app_config = self.config_manager.get_app_config()
            output_config = app_config.get('output', {})
            
            # 出力フォルダ
            if 'folder' in output_config:
                self.output_folder_edit.setText(output_config['folder'])
                
            # 出力形式
            if 'format' in output_config:
                index = self.format_combo.findText(output_config['format'])
                if index >= 0:
                    self.format_combo.setCurrentIndex(index)
                    
            # その他の設定
            if 'include_stats' in output_config:
                self.include_stats_check.setChecked(output_config['include_stats'])
                
            if 'include_keywords' in output_config:
                self.include_keywords_check.setChecked(output_config['include_keywords'])
                
            if 'include_original' in output_config:
                self.include_original_check.setChecked(output_config['include_original'])
                
        except Exception as e:
            self.logger.log_error(f"設定読み込みエラー: {e}")
            
    def save_settings(self):
        """設定を保存"""
        try:
            app_config = self.config_manager.get_app_config()
            
            if 'output' not in app_config:
                app_config['output'] = {}
                
            output_config = app_config['output']
            output_config['folder'] = self.output_folder_edit.text()
            output_config['format'] = self.format_combo.currentText()
            output_config['include_stats'] = self.include_stats_check.isChecked()
            output_config['include_keywords'] = self.include_keywords_check.isChecked()
            output_config['include_original'] = self.include_original_check.isChecked()
            
            self.config_manager.save_app_config(app_config)
            
        except Exception as e:
            self.logger.log_error(f"設定保存エラー: {e}")
            
    def set_output_format(self, format_type: str):
        """出力形式を設定"""
        index = self.format_combo.findText(format_type)
        if index >= 0:
            self.format_combo.setCurrentIndex(index)
            
    def set_output_directory(self, directory: str):
        """出力ディレクトリを設定"""
        self.output_folder_edit.setText(directory)
        
    def get_output_format(self) -> str:
        """出力形式を取得"""
        return self.format_combo.currentText()
        
    def get_output_directory(self) -> str:
        """出力ディレクトリを取得"""
        return self.output_folder_edit.text()
        
    def run_export(self):
        """出力を実行（エイリアス）"""
        self.execute_output() 